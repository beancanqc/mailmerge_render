"""
Word Document Splitter - Processor Class
Handles splitting Word documents by ranges or individual pages
"""

import os
import tempfile
import zipfile
import shutil
from datetime import datetime
from pathlib import Path
import uuid
from typing import List, Dict, Any, Optional, Tuple

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import base64

class WordSplitter:
    """Handles splitting Word documents into smaller files"""
    
    def __init__(self, session_id=None, output_folder=None):
        self.session_id = session_id or str(uuid.uuid4())
        self.document_path: Optional[str] = None
        self.document: Optional[Document] = None
        self.temp_dir = tempfile.mkdtemp()
        self.output_folder = output_folder or tempfile.mkdtemp()
        self.total_pages = 0
        self.page_breaks = []  # Track where page breaks occur
        
    def cleanup(self):
        """Clean up temporary files"""
        try:
            if self.document_path and os.path.exists(self.document_path):
                os.remove(self.document_path)
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
            print(f"✅ Cleanup completed for split session {self.session_id}")
        except Exception as e:
            print(f"❌ Cleanup error for session {self.session_id}: {e}")
    
    def load_document(self, document_path: str) -> bool:
        """Load and analyze Word document"""
        try:
            print(f"Loading document: {document_path}")
            
            # Validate file exists and is readable
            if not os.path.exists(document_path):
                print(f"Document file does not exist: {document_path}")
                return False
            
            # Load document
            self.document = Document(document_path)
            self.document_path = document_path
            
            # Analyze document structure
            self._analyze_document_structure()
            
            print(f"✅ Document loaded: {len(self.document.paragraphs)} paragraphs, estimated {self.total_pages} pages")
            return True
            
        except Exception as e:
            print(f"❌ Error loading document: {e}")
            return False
    
    def _analyze_document_structure(self):
        """Analyze document to estimate page breaks and structure"""
        try:
            # Estimate page breaks based on content
            # This is a simplified estimation - Word's pagination is complex
            paragraph_count = 0
            estimated_pages = 1
            
            for paragraph in self.document.paragraphs:
                paragraph_count += 1
                
                # Check for explicit page breaks
                for run in paragraph.runs:
                    if run._element.xml.find('w:br') is not None:
                        # Check if it's a page break
                        br_elements = run._element.findall('.//w:br', run._element.nsmap)
                        for br in br_elements:
                            if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                                estimated_pages += 1
                                self.page_breaks.append(paragraph_count)
                
                # Rough estimation: every 35-40 paragraphs = 1 page (very approximate)
                if paragraph_count % 40 == 0:
                    estimated_pages += 1
            
            self.total_pages = max(estimated_pages, 1)
            
            # If no explicit page breaks found, create estimated breaks
            if not self.page_breaks and self.total_pages > 1:
                paragraphs_per_page = max(len(self.document.paragraphs) // self.total_pages, 1)
                for page in range(1, self.total_pages):
                    self.page_breaks.append(page * paragraphs_per_page)
            
            print(f"Document analysis: {self.total_pages} estimated pages, {len(self.page_breaks)} page breaks detected")
            
        except Exception as e:
            print(f"Warning: Error analyzing document structure: {e}")
            self.total_pages = 1
            self.page_breaks = []
    
    def get_page_thumbnails(self) -> List[Dict[str, Any]]:
        """Generate page information for thumbnails (simplified for web)"""
        try:
            pages = []
            
            # For web implementation, we'll provide page ranges instead of actual thumbnails
            # Real thumbnail generation would require more complex libraries
            
            if self.total_pages <= 1:
                pages.append({
                    'page_number': 1,
                    'start_paragraph': 0,
                    'end_paragraph': len(self.document.paragraphs) - 1,
                    'content_preview': self._get_page_preview(0, len(self.document.paragraphs)),
                    'filename': os.path.basename(self.document_path) if self.document_path else 'document.docx'
                })
            else:
                # Create page ranges based on estimated breaks
                prev_break = 0
                for i, page_num in enumerate(range(1, self.total_pages + 1)):
                    if i < len(self.page_breaks):
                        current_break = self.page_breaks[i]
                    else:
                        current_break = len(self.document.paragraphs)
                    
                    pages.append({
                        'page_number': page_num,
                        'start_paragraph': prev_break,
                        'end_paragraph': min(current_break - 1, len(self.document.paragraphs) - 1),
                        'content_preview': self._get_page_preview(prev_break, current_break),
                        'filename': os.path.basename(self.document_path) if self.document_path else 'document.docx'
                    })
                    
                    prev_break = current_break
            
            return pages
            
        except Exception as e:
            print(f"Error generating page thumbnails: {e}")
            # Return single page as fallback
            return [{
                'page_number': 1,
                'start_paragraph': 0,
                'end_paragraph': len(self.document.paragraphs) - 1 if self.document else 0,
                'content_preview': 'Document preview unavailable',
                'filename': os.path.basename(self.document_path) if self.document_path else 'document.docx'
            }]
    
    def _get_page_preview(self, start_para: int, end_para: int) -> str:
        """Get text preview for a page range"""
        try:
            if not self.document:
                return ""
            
            preview_text = ""
            para_count = 0
            
            for i in range(start_para, min(end_para, len(self.document.paragraphs))):
                if para_count >= 3:  # Limit preview to first 3 paragraphs
                    preview_text += "..."
                    break
                    
                para_text = self.document.paragraphs[i].text.strip()
                if para_text:
                    preview_text += para_text[:100]  # Limit to 100 chars per paragraph
                    preview_text += " "
                    para_count += 1
            
            return preview_text.strip()[:200]  # Limit total preview
            
        except Exception as e:
            print(f"Error getting page preview: {e}")
            return "Preview unavailable"
    
    def split_by_range(self, ranges: List[Dict[str, int]], output_type: str) -> str:
        """Split document by specified page ranges"""
        try:
            print(f"Splitting document by ranges: {ranges}")
            
            if output_type == "separate":
                return self._split_ranges_separate(ranges)
            else:
                return self._split_ranges_merged(ranges)
                
        except Exception as e:
            print(f"Error splitting by range: {e}")
            raise e
    
    def _split_ranges_separate(self, ranges: List[Dict[str, int]]) -> str:
        """Split ranges into separate files and return ZIP path"""
        zip_path = os.path.join(self.output_folder, f"split_ranges_{self.session_id}.zip")
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for i, range_info in enumerate(ranges):
                start_page = range_info['start']
                end_page = range_info['end']
                
                # Create document for this range
                range_doc = self._extract_page_range(start_page, end_page)
                
                # Save to temporary file
                range_filename = f"pages_{start_page}-{end_page}.docx"
                range_path = os.path.join(self.temp_dir, range_filename)
                range_doc.save(range_path)
                
                # Add to ZIP
                zipf.write(range_path, range_filename)
                
                # Clean up temp file
                os.remove(range_path)
        
        return zip_path
    
    def _split_ranges_merged(self, ranges: List[Dict[str, int]]) -> str:
        """Merge ranges into single file"""
        merged_doc = Document()
        
        for i, range_info in enumerate(ranges):
            start_page = range_info['start']
            end_page = range_info['end']
            
            # Extract content for this range
            range_content = self._get_range_content(start_page, end_page)
            
            # Add content to merged document
            for paragraph in range_content:
                new_paragraph = merged_doc.add_paragraph()
                new_paragraph.text = paragraph.text
                
                # Copy paragraph formatting (simplified)
                new_paragraph.style = paragraph.style
            
            # Add page break between ranges (except for last range)
            if i < len(ranges) - 1:
                merged_doc.add_page_break()
        
        # Save merged document
        merged_path = os.path.join(self.output_folder, f"merged_ranges_{self.session_id}.docx")
        merged_doc.save(merged_path)
        
        return merged_path
    
    def split_by_pages(self, selected_pages: List[int], output_type: str) -> str:
        """Split document by individual pages"""
        try:
            print(f"Splitting document by pages: {selected_pages}")
            
            if output_type == "separate":
                return self._split_pages_separate(selected_pages)
            else:
                return self._split_pages_merged(selected_pages)
                
        except Exception as e:
            print(f"Error splitting by pages: {e}")
            raise e
    
    def _split_pages_separate(self, pages: List[int]) -> str:
        """Split pages into separate files and return ZIP path"""
        zip_path = os.path.join(self.output_folder, f"split_pages_{self.session_id}.zip")
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for page_num in sorted(pages):
                # Create document for this page
                page_doc = self._extract_single_page(page_num)
                
                # Save to temporary file
                page_filename = f"page_{page_num}.docx"
                page_path = os.path.join(self.temp_dir, page_filename)
                page_doc.save(page_path)
                
                # Add to ZIP
                zipf.write(page_path, page_filename)
                
                # Clean up temp file
                os.remove(page_path)
        
        return zip_path
    
    def _split_pages_merged(self, pages: List[int]) -> str:
        """Merge selected pages into single file"""
        merged_doc = Document()
        
        for i, page_num in enumerate(sorted(pages)):
            # Extract content for this page
            page_content = self._get_single_page_content(page_num)
            
            # Add content to merged document
            for paragraph in page_content:
                new_paragraph = merged_doc.add_paragraph()
                new_paragraph.text = paragraph.text
                
                # Copy paragraph formatting (simplified)
                new_paragraph.style = paragraph.style
            
            # Add page break between pages (except for last page)
            if i < len(pages) - 1:
                merged_doc.add_page_break()
        
        # Save merged document
        merged_path = os.path.join(self.output_folder, f"merged_pages_{self.session_id}.docx")
        merged_doc.save(merged_path)
        
        return merged_path
    
    def _extract_page_range(self, start_page: int, end_page: int) -> Document:
        """Extract a range of pages into a new document"""
        new_doc = Document()
        
        # Get content for the page range
        content = self._get_range_content(start_page, end_page)
        
        # Add content to new document
        for paragraph in content:
            new_paragraph = new_doc.add_paragraph()
            new_paragraph.text = paragraph.text
            
            # Copy basic formatting
            try:
                new_paragraph.style = paragraph.style
            except:
                pass  # Skip if style copying fails
        
        return new_doc
    
    def _extract_single_page(self, page_num: int) -> Document:
        """Extract a single page into a new document"""
        return self._extract_page_range(page_num, page_num)
    
    def _get_range_content(self, start_page: int, end_page: int) -> List[Any]:
        """Get paragraph content for a page range"""
        try:
            if not self.document:
                return []
            
            # Convert page numbers to paragraph indices
            start_para, end_para = self._pages_to_paragraphs(start_page, end_page)
            
            # Return paragraphs in range
            return self.document.paragraphs[start_para:end_para + 1]
            
        except Exception as e:
            print(f"Error getting range content: {e}")
            return []
    
    def _get_single_page_content(self, page_num: int) -> List[Any]:
        """Get paragraph content for a single page"""
        return self._get_range_content(page_num, page_num)
    
    def _pages_to_paragraphs(self, start_page: int, end_page: int) -> Tuple[int, int]:
        """Convert page numbers to paragraph indices"""
        try:
            # Handle single page document
            if self.total_pages <= 1:
                return 0, len(self.document.paragraphs) - 1
            
            # Calculate paragraph ranges based on page breaks
            start_para = 0
            end_para = len(self.document.paragraphs) - 1
            
            # Find start paragraph for start_page
            if start_page > 1:
                page_break_index = min(start_page - 2, len(self.page_breaks) - 1)
                if page_break_index >= 0:
                    start_para = self.page_breaks[page_break_index]
            
            # Find end paragraph for end_page
            if end_page < self.total_pages:
                page_break_index = min(end_page - 1, len(self.page_breaks) - 1)
                if page_break_index >= 0:
                    end_para = self.page_breaks[page_break_index] - 1
            
            # Ensure valid ranges
            start_para = max(0, start_para)
            end_para = min(len(self.document.paragraphs) - 1, end_para)
            end_para = max(start_para, end_para)  # Ensure end >= start
            
            return start_para, end_para
            
        except Exception as e:
            print(f"Error converting pages to paragraphs: {e}")
            return 0, len(self.document.paragraphs) - 1 if self.document else (0, 0)