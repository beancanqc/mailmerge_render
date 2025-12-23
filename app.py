"""
Mail Merge SaaS - Main Application
Flask web server for Render deployment - Word documents only
"""

import os
import tempfile
import zipfile
import shutil
from datetime import datetime
from pathlib import Path
import uuid

from flask import Flask, request, jsonify, send_file, session
from werkzeug.utils import secure_filename
from docx import Document
from docx.enum.text import WD_BREAK
import openpyxl
import re
from typing import List, Dict, Any, Optional

from jinja2 import Template

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Configure session management for production
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key-' + str(uuid.uuid4()))
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_USE_SIGNER'] = True
app.config['SESSION_KEY_PREFIX'] = 'mailmerge:'

# Configure folders
UPLOAD_FOLDER = tempfile.mkdtemp()
OUTPUT_FOLDER = tempfile.mkdtemp()
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_TEMPLATE_EXTENSIONS = {'docx'}
ALLOWED_DATA_EXTENSIONS = {'xlsx'}

def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

class MailMergeProcessor:
    def __init__(self, session_id=None):
        self.session_id = session_id or str(uuid.uuid4())
        self.template_path: Optional[str] = None
        self.data_path: Optional[str] = None
        self.data: List[Dict[str, Any]] = []
        self.headers: List[str] = []  # Store Excel column headers
        
    def cleanup(self):
        """Clean up temporary files"""
        try:
            if self.template_path and os.path.exists(self.template_path):
                os.remove(self.template_path)
                print(f"Cleaned up template: {self.template_path}")
            
            if self.data_path and os.path.exists(self.data_path):
                os.remove(self.data_path)
                print(f"Cleaned up data file: {self.data_path}")
                
        except Exception as e:
            print(f"Cleanup error: {str(e)}")
        
        # Reset state
        self.template_path = None
        self.data_path = None
        self.data = []
        self.headers = []
        
    def load_template(self, template_path: str) -> bool:
        """Load and validate Word template file"""
        try:
            # Clean up previous template
            if self.template_path and os.path.exists(self.template_path):
                os.remove(self.template_path)
            
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template file not found: {template_path}")
            
            if not template_path.lower().endswith('.docx'):
                raise ValueError("Template must be a Word document (.docx)")
            
            # Test if file can be opened
            doc = Document(template_path)
            doc = None  # Close the document
            
            self.template_path = template_path
            print(f"Template loaded successfully: {template_path}")
            return True
            
        except Exception as e:
            print(f"Error loading template: {str(e)}")
            return False
    
    def load_data(self, data_path: str) -> bool:
        """Load and validate Excel data file"""
        try:
            # Clean up previous data file
            if self.data_path and os.path.exists(self.data_path):
                os.remove(self.data_path)
            
            if not os.path.exists(data_path):
                raise FileNotFoundError(f"Data file not found: {data_path}")
            
            if not data_path.lower().endswith('.xlsx'):
                raise ValueError("Data file must be an Excel file (.xlsx)")
            
            # Load Excel data using openpyxl
            workbook = openpyxl.load_workbook(data_path, data_only=True)
            sheet = workbook.active
            
            if sheet.max_row <= 1:
                workbook.close()
                raise ValueError("Excel file appears to be empty or has no data")
            
            # Convert to list of dictionaries
            self.headers = []  # Store headers as instance variable
            for cell in sheet[1]:
                self.headers.append(str(cell.value) if cell.value is not None else "")
            
            self.data = []
            for row in sheet.iter_rows(min_row=2, values_only=True):
                row_data = {}
                for i, value in enumerate(row):
                    if i < len(self.headers):
                        row_data[self.headers[i]] = str(value) if value is not None else ""
                self.data.append(row_data)
            
            workbook.close()
            
            if not self.data:
                raise ValueError("No data rows found in Excel file")
            
            self.data_path = data_path
            print(f"Data loaded successfully: {len(self.data)} records from {data_path}")
            return True
            
        except Exception as e:
            print(f"Error loading data: {str(e)}")
            return False

    def _find_run_for_position(self, paragraph, text_position):
        """Find which run contains the text at the given position and return its formatting"""
        current_pos = 0
        
        for run in paragraph.runs:
            run_len = len(run.text)
            
            if current_pos <= text_position < current_pos + run_len:
                # Extract formatting from this run
                formatting = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb if run.font.color.rgb else None
                }
                return {'formatting': formatting}
            
            current_pos += run_len
        
        # If position not found, return formatting from first run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            formatting = {
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline,
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'font_color': first_run.font.color.rgb if first_run.font.color.rgb else None
            }
            return {'formatting': formatting}
        
        return None

    def _apply_formatting(self, run, formatting):
        """Apply formatting to a run"""
        try:
            if formatting.get('bold') is not None:
                run.bold = formatting['bold']
            if formatting.get('italic') is not None:
                run.italic = formatting['italic']
            if formatting.get('underline') is not None:
                run.underline = formatting['underline']
            if formatting.get('font_name'):
                run.font.name = formatting['font_name']
            if formatting.get('font_size'):
                run.font.size = formatting['font_size']
            if formatting.get('font_color'):
                run.font.color.rgb = formatting['font_color']
        except Exception as e:
            print(f"Error applying formatting: {e}")
            # Continue without formatting if there's an error

    def replace_merge_fields_advanced(self, paragraph, data_row: Dict[str, Any]):
        """Advanced merge field replacement that preserves individual character formatting"""
        full_text = paragraph.text
        
        # Find all merge fields
        merge_fields = re.finditer(r'\{\{(\w+)\}\}', full_text)
        merge_list = list(merge_fields)
        
        if not merge_list:
            return
        
        # Create a list to store new runs
        new_runs_data = []
        current_pos = 0
        
        for match in merge_list:
            field_name = match.group(1)
            start_pos = match.start()
            end_pos = match.end()
            replacement_text = str(data_row.get(field_name, ""))
            
            # Add text before the merge field
            if start_pos > current_pos:
                before_text = full_text[current_pos:start_pos]
                if before_text:
                    # Find the run that contains this text and its formatting
                    run_info = self._find_run_for_position(paragraph, current_pos)
                    new_runs_data.append({
                        'text': before_text,
                        'formatting': run_info['formatting'] if run_info else None
                    })
            
            # Add the replacement text with the formatting of the merge field location
            if replacement_text:
                run_info = self._find_run_for_position(paragraph, start_pos)
                new_runs_data.append({
                    'text': replacement_text,
                    'formatting': run_info['formatting'] if run_info else None
                })
            
            current_pos = end_pos
        
        # Add remaining text after the last merge field
        if current_pos < len(full_text):
            remaining_text = full_text[current_pos:]
            if remaining_text:
                run_info = self._find_run_for_position(paragraph, current_pos)
                new_runs_data.append({
                    'text': remaining_text,
                    'formatting': run_info['formatting'] if run_info else None
                })
        
        # Clear existing runs
        for run in paragraph.runs:
            run.clear()
        
        # Create new runs with preserved formatting
        for run_data in new_runs_data:
            if run_data['text']:
                new_run = paragraph.add_run(run_data['text'])
                if run_data['formatting']:
                    self._apply_formatting(new_run, run_data['formatting'])
    
    def replace_merge_fields(self, doc: Document, data_row: Dict[str, Any]) -> Document:
        """Replace merge fields with actual data while preserving formatting"""
        
        # Replace in paragraphs with advanced formatting preservation
        for paragraph in doc.paragraphs:
            if '{{' in paragraph.text and '}}' in paragraph.text:
                self.replace_merge_fields_advanced(paragraph, data_row)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{' in paragraph.text and '}}' in paragraph.text:
                            self.replace_merge_fields_advanced(paragraph, data_row)
        
        # Replace in headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    if '{{' in paragraph.text and '}}' in paragraph.text:
                        self.replace_merge_fields_advanced(paragraph, data_row)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if '{{' in paragraph.text and '}}' in paragraph.text:
                        self.replace_merge_fields_advanced(paragraph, data_row)
        
        return doc
    
    def generate_single_word(self, output_path: str) -> bool:
        """Generate a single Word document using SECTION BREAKS - Most reliable method"""
        try:
            if not self.template_path or not self.data:
                raise ValueError("Template and data must be loaded first")
            
            print(f"Creating single Word document with {len(self.data)} records using SECTION BREAKS...")
            
            # Start with first record - load fresh template and process
            final_doc = Document(self.template_path)
            final_doc = self.replace_merge_fields(final_doc, self.data[0])
            print(f"Added record 1 of {len(self.data)}")
            
            # Add remaining records using section breaks
            for i, row_data in enumerate(self.data[1:], 1):
                print(f"Adding record {i+1} of {len(self.data)} with section break...")
                
                # Load and process template for this record
                template_doc = Document(self.template_path)
                processed_doc = self.replace_merge_fields(template_doc, row_data)
                
                # Add new section with NEW_PAGE start (more reliable than page breaks)
                from docx.enum.section import WD_SECTION_START
                new_section = final_doc.add_section(WD_SECTION_START.NEW_PAGE)
                
                # Copy all content from processed template to the new section
                # Get the section element and body
                section_element = new_section._sectPr
                body_element = final_doc._body._body
                
                # Copy all paragraphs from processed document
                for para in processed_doc.paragraphs:
                    new_para = final_doc.add_paragraph()
                    
                    # Copy paragraph-level formatting
                    try:
                        new_para.style = para.style
                        new_para.alignment = para.alignment
                    except:
                        pass
                    
                    # Copy all runs with their formatting
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        
                        # Copy comprehensive formatting
                        try:
                            if run.bold is not None:
                                new_run.bold = run.bold
                            if run.italic is not None:
                                new_run.italic = run.italic
                            if run.underline is not None:
                                new_run.underline = run.underline
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb
                        except:
                            pass
                
                # Copy tables from processed document
                for table in processed_doc.tables:
                    new_table = final_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    
                    # Copy table content with formatting
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            new_cell = new_table.cell(row_idx, col_idx)
                            new_cell.text = ""  # Clear default text
                            
                            for para in cell.paragraphs:
                                if para.text.strip() or len(para.runs) > 0:
                                    new_para = new_cell.add_paragraph()
                                    for run in para.runs:
                                        new_run = new_para.add_run(run.text)
                                        try:
                                            if run.bold is not None:
                                                new_run.bold = run.bold
                                            if run.italic is not None:
                                                new_run.italic = run.italic
                                        except:
                                            pass
            
            # Save the final document
            final_doc.save(output_path)
            print(f"✅ Successfully created single Word document using section breaks")
            return True
            
        except Exception as e:
            print(f"❌ Error creating single Word document: {str(e)}")
            import traceback
            traceback.print_exc()
            
            # Fallback to traditional approach if section breaks fail
            print("🔄 Trying fallback approach with traditional page breaks...")
            return self.generate_single_word_fallback(output_path)

    def generate_single_word_fallback(self, output_path: str) -> bool:
        """Fallback method using traditional page breaks with XML manipulation"""
        try:
            print("Using fallback method with XML-level page break insertion...")
            
            # Process all records first
            all_processed_docs = []
            for i, row_data in enumerate(self.data):
                template_doc = Document(self.template_path)
                processed_doc = self.replace_merge_fields(template_doc, row_data)
                all_processed_docs.append(processed_doc)
                print(f"Processed record {i+1} of {len(self.data)}")
            
            # Start with first document
            final_doc = all_processed_docs[0]
            
            # Add remaining documents with XML page breaks
            for i, doc in enumerate(all_processed_docs[1:], 1):
                print(f"Merging record {i+1} with XML page break...")
                
                # Insert page break at XML level (more reliable)
                body = final_doc._body._body
                
                # Create page break paragraph
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls, qn
                
                # Add page break using XML
                page_break_xml = f'''
                <w:p {nsdecls('w')}>
                    <w:r>
                        <w:br w:type="page"/>
                    </w:r>
                </w:p>
                '''
                page_break_p = parse_xml(page_break_xml)
                body.append(page_break_p)
                
                # Copy all content from the document
                for para in doc.paragraphs:
                    new_para = final_doc.add_paragraph()
                    
                    try:
                        new_para.style = para.style
                        new_para.alignment = para.alignment
                    except:
                        pass
                    
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        try:
                            if run.bold is not None:
                                new_run.bold = run.bold
                            if run.italic is not None:
                                new_run.italic = run.italic
                            if run.underline is not None:
                                new_run.underline = run.underline
                        except:
                            pass
                
                # Copy tables
                for table in doc.tables:
                    new_table = final_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            new_table.cell(row_idx, col_idx).text = cell.text
            
            final_doc.save(output_path)
            print("✅ Fallback method successful")
            return True
            
        except Exception as e:
            print(f"❌ Fallback method also failed: {str(e)}")
            return False
    
    def generate_multiple_word(self, output_dir: str) -> bool:
        """Generate multiple Word documents (one per record) - filename based on first column"""
        try:
            if not self.template_path or not self.data:
                raise ValueError("Template and data must be loaded first")
            
            os.makedirs(output_dir, exist_ok=True)
            
            # Get the first column header for filename generation
            first_column_header = self.headers[0] if self.headers else None
            print(f"Using first column '{first_column_header}' for filenames")
            
            for index, row_data in enumerate(self.data):
                # Load template
                doc = Document(self.template_path)
                
                # Replace merge fields
                processed_doc = self.replace_merge_fields(doc, row_data)
                
                # Generate filename using first column value
                if first_column_header and first_column_header in row_data:
                    filename_value = row_data[first_column_header]
                else:
                    # Fallback to first value in row or record number
                    filename_value = list(row_data.values())[0] if row_data else f"record_{index+1}"
                
                # Clean filename - remove invalid characters for file system
                safe_filename = re.sub(r'[<>:"/\\|?*]', '_', str(filename_value))
                # Also remove leading/trailing spaces and dots
                safe_filename = safe_filename.strip('. ')
                # Ensure filename is not empty
                if not safe_filename:
                    safe_filename = f"record_{index+1}"
                
                output_path = os.path.join(output_dir, f"{safe_filename}.docx")
                
                # Handle duplicate filenames by adding index
                counter = 1
                base_path = output_path
                while os.path.exists(output_path):
                    name_without_ext = os.path.splitext(base_path)[0]
                    output_path = f"{name_without_ext}_{counter}.docx"
                    counter += 1
                
                processed_doc.save(output_path)
                print(f"Created: {os.path.basename(output_path)}")
            
            return True
            
        except Exception as e:
            print(f"Error creating multiple Word files: {str(e)}")
            return False
    

    





    
    def process_merge(self, output_format: str, output_path: str) -> bool:
        """Main processing function"""
        try:
            if not self.template_path or not self.data:
                raise ValueError("Both template and data files must be loaded")
            
            # Process based on output format
            if output_format == "single-word":
                return self.generate_single_word(output_path)
            elif output_format == "multiple-word":
                return self.generate_multiple_word(output_path)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")
                
        except Exception as e:
            print(f"Error processing mail merge: {str(e)}")
            return False

# Store processors per session
processors = {}

def get_processor():
    """Get or create processor for current session"""
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
        print(f"🆕 Created new session: {session['session_id']}")
    
    session_id = session['session_id']
    print(f"🔄 Using session: {session_id}")
    
    if session_id not in processors:
        processors[session_id] = MailMergeProcessor(session_id)
        print(f"🆕 Created new processor for session: {session_id}")
    else:
        print(f"♻️  Reusing existing processor for session: {session_id}")
    
    print(f"📊 Total active processors: {len(processors)}")
    return processors[session_id]

def cleanup_old_processors():
    """Clean up old processors (simple cleanup)"""
    if len(processors) > 50:  # Clean up if too many processors
        old_sessions = list(processors.keys())[:25]
        for session_id in old_sessions:
            processors[session_id].cleanup()
            del processors[session_id]

# Flask Routes
@app.route('/')
def index():
    """Serve the main page"""
    try:
        with open('index.html', 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return "<h1>Mail Merge SaaS</h1><p>Main page not found. Please upload index.html</p>"

@app.route('/mailmerge')
def mailmerge():
    """Serve the mail merge page"""
    try:
        with open('mailmerge.html', 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return "<h1>Mail Merge</h1><p>Mail merge page not found. Please upload mailmerge.html</p>"

@app.route('/style.css')
def serve_css():
    """Serve CSS file"""
    try:
        with open('style.css', 'r', encoding='utf-8') as f:
            css_content = f.read()
        response = app.response_class(
            response=css_content,
            status=200,
            mimetype='text/css'
        )
        return response
    except FileNotFoundError:
        return "/* CSS file not found */", 404

@app.route('/mailmerge.js')
def serve_js():
    """Serve JavaScript file"""
    try:
        with open('mailmerge.js', 'r', encoding='utf-8') as f:
            js_content = f.read()
        response = app.response_class(
            response=js_content,
            status=200,
            mimetype='application/javascript'
        )
        return response
    except FileNotFoundError:
        return "/* JavaScript file not found */", 404

@app.route('/static/<filename>')
def serve_static(filename):
    """Serve static files"""
    try:
        static_folder = 'static'
        file_path = os.path.join(static_folder, filename)
        if os.path.exists(file_path):
            return send_file(file_path)
        else:
            return "File not found", 404
    except Exception as e:
        return f"Error: {str(e)}", 500

@app.route('/upload_template', methods=['POST'])
def upload_template():
    """Handle template file upload"""
    try:
        print("Template upload request received")
        cleanup_old_processors()
        
        processor = get_processor()
        
        if 'file' not in request.files:
            print("No file in request")
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        file = request.files['file']
        if file.filename == '':
            print("Empty filename")
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        print(f"Template file: {file.filename}")
        
        if not allowed_file(file.filename, ALLOWED_TEMPLATE_EXTENSIONS):
            print(f"Invalid file type: {file.filename}")
            return jsonify({'success': False, 'error': 'Invalid file type. Please upload a .docx file'}), 400
        
        # Create unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        filename = f"template_{processor.session_id}_{timestamp}_{secure_filename(file.filename)}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        print(f"Template saved to: {filepath}")
        
        # Load template
        if processor.load_template(filepath):
            print(f"✅ Template loaded successfully for session {processor.session_id}")
            print(f"   Template path set to: {processor.template_path}")
            print(f"   File exists: {os.path.exists(processor.template_path)}")
            return jsonify({
                'success': True,
                'message': f'Template uploaded successfully: {file.filename}',
                'filepath': filepath,
                'filename': file.filename
            })
        else:
            print(f"❌ Failed to load template for session {processor.session_id}")
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'error': 'Invalid template file'}), 400
            
    except Exception as e:
        print(f"Template upload error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/upload_data', methods=['POST'])
def upload_data():
    """Handle data file upload"""
    try:
        print("Data upload request received")
        cleanup_old_processors()
        
        processor = get_processor()
        
        if 'file' not in request.files:
            print("No file in request")
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        file = request.files['file']
        if file.filename == '':
            print("Empty filename")
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        print(f"Data file: {file.filename}")
        
        if not allowed_file(file.filename, ALLOWED_DATA_EXTENSIONS):
            print(f"Invalid file type: {file.filename}")
            return jsonify({'success': False, 'error': 'Invalid file type. Please upload an Excel file (.xlsx)'}), 400
        
        # Create unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        filename = f"data_{processor.session_id}_{timestamp}_{secure_filename(file.filename)}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        print(f"Data saved to: {filepath}")
        
        # Load data
        if processor.load_data(filepath):
            # Return preview of data
            preview_data = processor.data[:3]  # First 3 rows
            columns = list(processor.data[0].keys()) if processor.data else []
            total_rows = len(processor.data)
            
            return jsonify({
                'success': True,
                'message': f'Data uploaded successfully: {file.filename}',
                'filepath': filepath,
                'filename': file.filename,
                'preview': preview_data,
                'columns': columns,
                'total_rows': total_rows
            })
        else:
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'error': 'Invalid data file'}), 400
            
    except Exception as e:
        print(f"Data upload error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/check_status', methods=['GET'])
def check_status():
    """Check current upload status - with fallback file checking"""
    try:
        processor = get_processor()
        
        # Debug logging
        print(f"🔍 Status check for session: {processor.session_id}")
        print(f"   Template path: {processor.template_path}")
        print(f"   Template exists: {processor.template_path and os.path.exists(processor.template_path) if processor.template_path else False}")
        print(f"   Data loaded: {len(processor.data) if processor.data else 0} records")
        
        # Fallback: If processor doesn't have files, check for recent uploads
        template_loaded = processor.template_path is not None and os.path.exists(processor.template_path) if processor.template_path else False
        data_loaded = processor.data_path is not None and len(processor.data) > 0
        
        # If files not found in processor, check for recent uploads in the directory
        if not template_loaded or not data_loaded:
            print("🔍 Checking for recent uploads as fallback...")
            if os.path.exists(app.config['UPLOAD_FOLDER']):
                files = os.listdir(app.config['UPLOAD_FOLDER'])
                print(f"   Found files: {files}")
                
                # Look for recent template files
                if not template_loaded:
                    for file in files:
                        if file.startswith(f'template_{processor.session_id}') and file.endswith('.docx'):
                            template_path = os.path.join(app.config['UPLOAD_FOLDER'], file)
                            if processor.load_template(template_path):
                                template_loaded = True
                                print(f"✅ Recovered template: {file}")
                                break
                
                # Look for recent data files  
                if not data_loaded:
                    for file in files:
                        if file.startswith(f'data_{processor.session_id}') and file.endswith('.xlsx'):
                            data_path = os.path.join(app.config['UPLOAD_FOLDER'], file)
                            if processor.load_data(data_path):
                                data_loaded = True
                                print(f"✅ Recovered data: {file}")
                                break
        
        result = {
            'template_loaded': template_loaded,
            'data_loaded': data_loaded,
            'template_path': processor.template_path,
            'data_records': len(processor.data) if processor.data else 0,
            'session_id': processor.session_id  # Add for debugging
        }
        
        print(f"   Returning status: {result}")
        return jsonify(result)
        
    except Exception as e:
        print(f"❌ Status check error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/process_merge', methods=['POST'])
def process_merge():
    """Process the mail merge - Support both Word and PDF formats"""
    try:
        print("Process merge request received")
        
        processor = get_processor()
        data = request.get_json()
        output_format = data.get('format', 'single-word')
        
        print(f"Output format: {output_format}")
        print(f"Template loaded: {processor.template_path is not None}")
        print(f"Data loaded: {len(processor.data) if processor.data else 0} records")
        
        if not processor.template_path or not processor.data:
            error_msg = f"Missing files - Template: {processor.template_path is not None}, Data: {len(processor.data) if processor.data else 0} records"
            print(error_msg)
            return jsonify({'success': False, 'error': 'Please upload both template and data files first'}), 400
        
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        
        # Determine file extension and format type
        is_single = 'single' in output_format
        file_ext = '.docx'
        format_name = 'Word'
        
        if is_single:
            # Single file output
            output_filename = f"mailmerge_result_{processor.session_id}_{timestamp}{file_ext}"
            output_path = os.path.join(OUTPUT_FOLDER, output_filename)
            
            if processor.process_merge(output_format, output_path):
                return jsonify({
                    'success': True,
                    'message': f'Mail merge completed successfully! {format_name} formatting preserved.',
                    'download_url': f'/download/{output_filename}',
                    'filename': output_filename
                })
            else:
                return jsonify({'success': False, 'error': 'Failed to process mail merge'}), 500
                
        else:  # multiple files
            # Multiple files - create ZIP
            output_dir = os.path.join(OUTPUT_FOLDER, f"mailmerge_{processor.session_id}_{timestamp}")
            zip_filename = f"mailmerge_results_{processor.session_id}_{timestamp}.zip"
            zip_path = os.path.join(OUTPUT_FOLDER, zip_filename)
            
            if processor.process_merge(output_format, output_dir):
                # Create ZIP file
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    for root, dirs, files in os.walk(output_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, output_dir)
                            zipf.write(file_path, arcname)
                
                # Clean up individual files
                shutil.rmtree(output_dir)
                
                return jsonify({
                    'success': True,
                    'message': f'Mail merge completed! Generated {len(processor.data)} {format_name} documents with preserved formatting.',
                    'download_url': f'/download/{zip_filename}',
                    'filename': zip_filename
                })
            else:
                return jsonify({'success': False, 'error': 'Failed to process mail merge'}), 500
                
    except Exception as e:
        print(f"Process merge error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Download processed file"""
    try:
        file_path = os.path.join(OUTPUT_FOLDER, filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/health')
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy', 
        'service': 'Mail Merge SaaS - Word & PDF Support',
        'active_sessions': len(processors),
        'upload_folder': app.config['UPLOAD_FOLDER'],
        'output_folder': OUTPUT_FOLDER
    })

@app.route('/debug')
def debug_info():
    """Debug endpoint to check application state"""
    try:
        processor = get_processor()
        
        # List files in upload directory
        upload_files = []
        if os.path.exists(app.config['UPLOAD_FOLDER']):
            upload_files = os.listdir(app.config['UPLOAD_FOLDER'])
        
        return jsonify({
            'session_id': processor.session_id,
            'session_data': dict(session),
            'processors_count': len(processors),
            'processor_template_path': processor.template_path,
            'processor_data_records': len(processor.data) if processor.data else 0,
            'upload_folder': app.config['UPLOAD_FOLDER'],
            'upload_files': upload_files,
            'template_file_exists': os.path.exists(processor.template_path) if processor.template_path else False
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # Development server
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)