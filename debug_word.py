#!/usr/bin/env python3
"""
Debug script to understand pagination in Word documents
"""

import os
from docx import Document

def debug_word_document(file_path):
    """Debug a Word document to understand its structure"""
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    doc = Document(file_path)
    
    print(f"Analyzing: {file_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    print()
    
    # Look for page breaks
    page_breaks = 0
    for i, paragraph in enumerate(doc.paragraphs):
        for run in paragraph.runs:
            if '\f' in run.text or '\x0c' in run.text:
                page_breaks += 1
                print(f"Page break found in paragraph {i}: '{paragraph.text}'")
        
        # Check for page break elements
        for element in paragraph._element:
            if hasattr(element, 'tag') and 'pageBreak' in str(element.tag):
                page_breaks += 1
                print(f"Page break element found in paragraph {i}")
    
    print(f"Total page breaks found: {page_breaks}")
    print(f"Expected pages: {page_breaks + 1}")
    
    # Show paragraph content summary
    print("\nParagraph summary:")
    for i, paragraph in enumerate(doc.paragraphs[:10]):  # First 10 paragraphs
        text = paragraph.text.strip()
        if text:
            print(f"Para {i}: {text[:50]}...")
        else:
            print(f"Para {i}: (empty)")

if __name__ == "__main__":
    # Look for Word documents in current directory
    word_files = [f for f in os.listdir('.') if f.endswith('.docx') and not f.startswith('~')]
    
    if word_files:
        print("Found Word documents:")
        for i, file in enumerate(word_files):
            print(f"{i+1}. {file}")
        
        # Debug the first one or let user choose
        if len(word_files) == 1:
            debug_word_document(word_files[0])
        else:
            print(f"\nDebugging first file: {word_files[0]}")
            debug_word_document(word_files[0])
    else:
        print("No Word documents found in current directory")