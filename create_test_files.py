import os
os.chdir(r'c:\Users\julie\Desktop\site 1 downloaded from github 13 décembre 2025 - Copie')

from docx import Document
import openpyxl

# Create test template
print("Creating test template...")
doc = Document()
doc.add_heading('Mail Merge Test Template', 0)
p = doc.add_paragraph('Dear ')
p.add_run('{{first_name}}').bold = True
p.add_run(' ')
p.add_run('{{last_name}}').bold = True
p.add_run(',')
doc.add_paragraph('')
doc.add_paragraph('Welcome to {{company}}! Your email is: {{email}}')
doc.add_paragraph('')
doc.add_paragraph('Best regards,')
doc.add_paragraph('The Team')
doc.save('test_template.docx')
print('✅ Created test_template.docx')

# Create test data
print("Creating test data...")
wb = openpyxl.Workbook()
ws = wb.active
ws['A1'] = 'first_name'
ws['B1'] = 'last_name'
ws['C1'] = 'company'
ws['D1'] = 'email'
ws['A2'] = 'John'
ws['B2'] = 'Doe'
ws['C2'] = 'ABC Corp'
ws['D2'] = 'john@abc.com'
ws['A3'] = 'Jane'
ws['B3'] = 'Smith'
ws['C3'] = 'XYZ Ltd'
ws['D3'] = 'jane@xyz.com'
wb.save('test_data.xlsx')
print('✅ Created test_data.xlsx')