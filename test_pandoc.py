#!/usr/bin/env python3
"""
Test script for Pandoc DOCX to PDF conversion
Run this after deploying to Render to verify Pandoc installation
"""

import subprocess
import tempfile
import os
from docx import Document

def test_pandoc_installation():
    """Test if Pandoc is properly installed"""
    try:
        result = subprocess.run(['pandoc', '--version'], capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            print("✅ Pandoc is installed!")
            print(f"Version: {result.stdout.split()[1]}")
            return True
        else:
            print("❌ Pandoc not working properly")
            return False
    except FileNotFoundError:
        print("❌ Pandoc not found - check aptfile installation")
        return False
    except Exception as e:
        print(f"❌ Error checking Pandoc: {e}")
        return False

def create_test_docx():
    """Create a simple test DOCX file with merge fields"""
    doc = Document()
    doc.add_heading('Mail Merge Test Document', 0)
    
    para1 = doc.add_paragraph('Dear ')
    para1.add_run('{{FirstName}} {{LastName}}').bold = True
    para1.add_run(',')
    
    doc.add_paragraph('This is a test document to verify that Pandoc can successfully convert DOCX files to PDF format.')
    
    para2 = doc.add_paragraph('Your email address is: ')
    para2.add_run('{{Email}}').italic = True
    
    doc.add_paragraph('If you can see this content in a PDF file, the conversion is working correctly!')
    
    # Save to temporary file
    temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(temp_docx.name)
    temp_docx.close()
    
    print(f"✅ Created test DOCX: {temp_docx.name}")
    return temp_docx.name

def test_pandoc_conversion(docx_path):
    """Test Pandoc DOCX to PDF conversion"""
    try:
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        # Test wkhtmltopdf engine first
        cmd = [
            'pandoc',
            docx_path,
            '-o', pdf_path,
            '--pdf-engine=wkhtmltopdf',
            '--pdf-engine-opt=--enable-local-file-access',
            '--pdf-engine-opt=--page-size', 'A4'
        ]
        
        print("🔄 Testing Pandoc with wkhtmltopdf...")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0 and os.path.exists(pdf_path):
            print("✅ wkhtmltopdf conversion successful!")
            print(f"PDF created: {pdf_path}")
            return True
        else:
            print("❌ wkhtmltopdf failed, trying LaTeX...")
            print(f"Error: {result.stderr}")
            
            # Try LaTeX fallback
            cmd_latex = ['pandoc', docx_path, '-o', pdf_path, '--pdf-engine=pdflatex']
            result2 = subprocess.run(cmd_latex, capture_output=True, text=True, timeout=90)
            
            if result2.returncode == 0 and os.path.exists(pdf_path):
                print("✅ LaTeX conversion successful!")
                print(f"PDF created: {pdf_path}")
                return True
            else:
                print("❌ Both conversion methods failed")
                print(f"LaTeX error: {result2.stderr}")
                return False
                
    except subprocess.TimeoutExpired:
        print("❌ Conversion timed out")
        return False
    except Exception as e:
        print(f"❌ Conversion error: {e}")
        return False

def main():
    """Run all tests"""
    print("🚀 Testing Pandoc PDF conversion setup...")
    print("=" * 50)
    
    # Test 1: Pandoc installation
    if not test_pandoc_installation():
        print("\n❌ Pandoc installation test failed!")
        return
    
    print("\n" + "=" * 50)
    
    # Test 2: Create test document
    try:
        docx_path = create_test_docx()
    except Exception as e:
        print(f"❌ Failed to create test document: {e}")
        return
    
    print("=" * 50)
    
    # Test 3: PDF conversion
    success = test_pandoc_conversion(docx_path)
    
    # Cleanup
    try:
        os.unlink(docx_path)
        pdf_path = docx_path.replace('.docx', '.pdf')
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
    except:
        pass
    
    print("\n" + "=" * 50)
    if success:
        print("🎉 ALL TESTS PASSED! Pandoc PDF conversion is ready!")
    else:
        print("❌ Tests failed - check Pandoc installation and dependencies")

if __name__ == "__main__":
    main()